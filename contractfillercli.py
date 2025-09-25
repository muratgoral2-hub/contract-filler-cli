import argparse
import csv
import json
import logging
import openpyxl
import unicodedata
from typing import Iterator, Iterable, Optional, Dict, List, Union, Any
from docx import Document
from pathlib import Path
from docx2pdf import convert
from datetime import datetime
from reportlab.pdfgen import canvas
from PyPDF2 import PdfReader, PdfWriter

# Global logger configuration
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)


# ---------------------------------------
# Helper: Normalize header names
# ---------------------------------------
def normalize_header(header: str) -> str:
    if not header:
        return ""
    # Normalize special characters (e.g., Turkish to ASCII)
    header = unicodedata.normalize("NFKD", header).encode("ascii", "ignore").decode("ascii")
    # Lowercase and replace spaces with underscores
    return header.strip().lower().replace(" ", "_")


# ---------------------------------------
# Date formatting helper
# ---------------------------------------
def format_date(value):
    try:
        if isinstance(value, str):
            return datetime.strptime(value, "%Y-%m-%d").strftime("%d/%m/%Y")
        elif isinstance(value, datetime):
            return value.strftime("%d/%m/%Y")
        else:
            return value
    except Exception:
        return str(value)


# ---------------------------------------
# Add logo to PDF (overlay on first page)
# ---------------------------------------
def add_logo_to_pdf(pdf_path, logo_path, output_path):
    reader = PdfReader(pdf_path)
    writer = PdfWriter()

    first_page = reader.pages[0]
    width = float(first_page.mediabox.width)
    height = float(first_page.mediabox.height)

    overlay_path = Path(output_path).with_name(Path(output_path).stem + "_overlay.pdf")
    c = canvas.Canvas(str(overlay_path), pagesize=(width, height))
    c.drawImage(logo_path, 50, height - 100, width=120, preserveAspectRatio=True, mask='auto')
    c.save()

    overlay_reader = PdfReader(str(overlay_path))
    overlay_page = overlay_reader.pages[0]
    first_page.merge_page(overlay_page)

    writer.add_page(first_page)
    for page in reader.pages[1:]:
        writer.add_page(page)

    with open(output_path, "wb") as f_out:
        writer.write(f_out)

    overlay_path.unlink(missing_ok=True)


# ---------------------------------------
# Read data from CSV, Excel, JSON, JSONL
# ---------------------------------------
def read_data(
    file_path: str,
    *,
    as_list: bool = False,
    return_dataframe: bool = False,
    normalize_headers: bool = True,
    column_map: Optional[Dict[str, str]] = None,
    required_fields: Optional[Iterable[str]] = None,
    date_fields: Optional[Iterable[str]] = None,
    date_formatter: Optional[Any] = None,
    csv_encoding: str = "utf-8",
    csv_delimiter: str = ",",
    invalid_sink_path: Optional[str] = "Errors/errors.json",
    progress: bool = True,
) -> Union[Iterator[Dict[str, Any]], List[Dict[str, Any]]]:

    def log(level: str, msg: str):
        if logger:
            getattr(logger, level, logger.info)(msg)
        else:
            print(f"[{level.upper()}] {msg}")

    path = Path(file_path)
    suffix = path.suffix.lower()
    colmap = {normalize_header(k): v for k, v in (column_map or {}).items()}
    req = tuple(normalize_header(f) for f in (required_fields or ()))
    date_keys = tuple(normalize_header(f) for f in (date_fields or ()))
    invalid_rows: list[dict] = []
    want_list = bool(as_list or return_dataframe)
    out_rows: list[dict] | None = [] if want_list else None

    iter_wrap = (lambda it: it)
    if progress:
        try:
            from tqdm import tqdm
            iter_wrap = lambda it: tqdm(it, desc=f"Reading {path.name}")
        except Exception:
            log("warning", "tqdm not available; continuing without progress bar.")

    def normalize_row(raw: dict) -> dict:
        row = {}
        for k, v in raw.items():
            new_key = normalize_header(str(k)) if k else ""
            if colmap:
                new_key = colmap.get(new_key, new_key)
            row[new_key] = v.strip() if isinstance(v, str) else v
        if date_formatter and date_keys:
            for dk in date_keys:
                if dk in row and row[dk]:
                    try:
                        if callable(date_formatter):
                            row[dk] = date_formatter(row[dk])
                        elif isinstance(date_formatter, dict) and row[dk] in date_formatter:
                            row[dk] = date_formatter[row[dk]]
                    except Exception as e:
                        log("warning", f"date_formatter {dk} error: {e}")
        return row

    def is_valid(row: dict) -> tuple[bool, list[str]]:
        missing = [f for f in req if not row.get(f)]
        return (len(missing) == 0, missing)

    def emit_row(row: dict):
        if want_list:
            out_rows.append(row)  # type: ignore[arg-type]
        else:
            return row

    def record_invalid_sink(row: dict, reason: list[str]):
        row_copy = dict(row)
        row_copy["_errors"] = reason
        invalid_rows.append(row_copy)

    def read_csv() -> Iterator[dict]:
        with path.open(newline="", encoding=csv_encoding) as f:
            rdr = csv.DictReader(f, delimiter=csv_delimiter)
            for raw in iter_wrap(rdr):
                if not raw or not any(raw.values()):
                    continue
                row = normalize_row(raw)
                ok, missing = is_valid(row)
                if not ok:
                    record_invalid_sink(row, [f"missing: {m}" for m in missing])
                    continue
                yielded = emit_row(row)
                if yielded is not None:
                    yield yielded

    def read_xlsx() -> Iterator[dict]:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        ws = wb.active
        headers = [("" if c.value is None else str(c.value)) for c in ws[1]]
        for row_vals in iter_wrap(ws.iter_rows(min_row=2, values_only=True)):
            if not row_vals or not any(row_vals):
                continue
            raw = dict(zip(headers, row_vals))
            row = normalize_row(raw)
            ok, missing = is_valid(row)
            if not ok:
                record_invalid_sink(row, [f"missing: {m}" for m in missing])
                continue
            yielded = emit_row(row)
            if yielded is not None:
                yield yielded

    def read_json() -> Iterator[dict]:
        with path.open(encoding="utf-8") as f:
            data = json.load(f)
        items = data if isinstance(data, list) else [data]
        for raw in iter_wrap(items):
            if not raw or (isinstance(raw, dict) and not any(raw.values())):
                continue
            if not isinstance(raw, dict):
                log("warning", f"JSON record is not a dict; skipped: {type(raw).__name__}")
                continue
            row = normalize_row(raw)
            ok, missing = is_valid(row)
            if not ok:
                record_invalid_sink(row, [f"missing: {m}" for m in missing])
                continue
            yielded = emit_row(row)
            if yielded is not None:
                yield yielded

    def read_jsonl() -> Iterator[dict]:
        with path.open(encoding="utf-8") as f:
            for line in iter_wrap(f):
                line = line.strip()
                if not line:
                    continue
                try:
                    raw = json.loads(line)
                except Exception as e:
                    record_invalid_sink({"_raw": line}, [f"json_decode:{e}"])
                    continue
                if not isinstance(raw, dict):
                    record_invalid_sink({"_raw": line}, ["json_not_object"])
                    continue
                row = normalize_row(raw)
                ok, missing = is_valid(row)
                if not ok:
                    record_invalid_sink(row, [f"missing: {m}" for m in missing])
                    continue
                yielded = emit_row(row)
                if yielded is not None:
                    yield yielded

    try:
        if suffix == ".csv":
            iterator = read_csv()
        elif suffix == ".xlsx":
            iterator = read_xlsx()
        elif suffix == ".json":
            iterator = read_json()
        elif suffix == ".jsonl":
            iterator = read_jsonl()
        else:
            raise ValueError("Unsupported file format. Use .csv, .xlsx, .json or .jsonl")

        if want_list:
            for _ in iterator:
                pass
        else:
            for item in iterator:
                yield item

    except FileNotFoundError:
        log("error", f"File not found: {file_path}")
    except PermissionError:
        log("error", f"Permission denied: {file_path}")
    except Exception as e:
        log("error", f"Unexpected error while reading {file_path}: {e}")

    if invalid_rows:
        sink = Path(invalid_sink_path) if invalid_sink_path else path.with_suffix(".invalid.json")
        sink.parent.mkdir(parents=True, exist_ok=True)
        try:
            sink.write_text(json.dumps(invalid_rows, ensure_ascii=False, indent=2), encoding="utf-8")
            log("warning", f"{len(invalid_rows)} invalid rows saved to {sink}")
        except Exception as e:
            log("error", f"Could not write invalid rows to {sink}: {e}")

    if return_dataframe:
        try:
            import pandas as pd
        except Exception as e:
            log("error", f"pandas not available for return_dataframe=True: {e}")
            return out_rows if as_list else []
        df = pd.DataFrame(out_rows)
        return df

    if as_list:
        return out_rows


# ---------------------------------------
# Build output filename from client data
# ---------------------------------------
def build_filename(data: dict, out_dir: str | Path) -> Path:
    name = (str(data.get("name", "noname")) or "noname").strip()
    surname = (str(data.get("surname", "nosurname")) or "nosurname").strip()
    return Path(out_dir) / f"{name}_{surname}.docx"


# ---------------------------------------
# Fill contract template with client data
# ---------------------------------------
def fill_contract(template_path, data_file, out_dir, logo_path=None):
    Path(out_dir).mkdir(parents=True, exist_ok=True)
    for data in read_data(data_file):
        doc = Document(template_path)

        # 🔹 Replace placeholders in paragraphs (keeping style)
        for p in doc.paragraphs:
            for key, value in data.items():
                placeholder = "{" + key + "}"
                if placeholder in p.text:
                    for run in p.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        # 🔹 Replace placeholders in tables (keeping style)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in data.items():
                            placeholder = "{" + key + "}"
                            if placeholder in p.text:
                                for run in p.runs:
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(value))

        out_docx = build_filename(data, out_dir)
        doc.save(str(out_docx))

        out_pdf = str(out_docx).replace(".docx", ".pdf")
        convert(str(out_docx), out_pdf)

        if logo_path:
            final_pdf = out_pdf.replace(".pdf", "_with_logo.pdf")
            add_logo_to_pdf(out_pdf, logo_path, final_pdf)
            Path(out_pdf).unlink(missing_ok=True)
            out_pdf = final_pdf

        print(f"[OK] Contract created for {data.get('name', '?')} {data.get('surname', '?')}: {out_pdf}")


# ---------------------------------------
# CLI entry point
# ---------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Contract Auto-filler CLI")
    parser.add_argument("--template", "-t", default="contract_template.docx", help="Path to template file")
    parser.add_argument("--data", "-d", default="client.xlsx", help="Path to data file")
    parser.add_argument("--out", "-o", default="contract", help="Path to output directory")
    parser.add_argument("--logo", "-l", default=None, help="Path to logo file")

    args = parser.parse_args()

    if not Path(args.template).exists():
        print(f"[ERROR] Template file not found: {args.template}")
        return

    if not Path(args.data).exists():
        print(f"[ERROR] Data file not found: {args.data}")
        return

    fill_contract(args.template, args.data, args.out, args.logo)


if __name__ == "__main__":
    main()
